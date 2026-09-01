"""
De un archivo a fragmentos consultables.

Tres pasos, y cada uno puede fallar de forma distinta, así que cada uno
deja constancia: extraer el texto, trocearlo, vectorizarlo.

La regla que ordena el módulo: **un documento a medio procesar no se
consulta**. Un PDF del que se extrajeron 3 de 40 páginas responde con
seguridad sobre las 3 y calla las 37, que es la peor forma de equivocarse.
Por eso el estado solo pasa a LISTO al final, y la búsqueda ignora todo lo
que no esté LISTO.
"""
from __future__ import annotations

import hashlib
import os
import re
from datetime import datetime, timezone
from pathlib import Path

import db
import ia
import tabla

ARCHIVOS = Path(os.getenv("ARCHIVOS", "./archivos"))

# Un fragmento debe ser lo bastante grande para que se entienda solo y lo
# bastante chico para que quepan varios en el contexto. 1.200 caracteres son
# unos dos párrafos: suficiente para una idea completa.
TAMANO = int(os.getenv("FRAGMENTO_TAMANO", "1200"))
SOLAPE = int(os.getenv("FRAGMENTO_SOLAPE", "150"))

# Las tablas de cifras se trocean más corto. La misma cantidad de
# caracteres cuesta el triple de tokens cuando son números -- cada cifra se
# parte en varios -- y lo que le cuesta tiempo al modelo son los tokens, no
# los caracteres. Medido en el servidor de oficina: ocho fragmentos de un
# balance en Excel tardaron 288 segundos solo en leerse.
TAMANO_TABLA = int(os.getenv("FRAGMENTO_TAMANO_TABLA", "500"))
TIPOS_TABLA = {"Excel", "CSV"}

TIPOS = {".pdf": "PDF", ".docx": "Word", ".xlsx": "Excel", ".xlsm": "Excel",
         ".txt": "Texto", ".md": "Markdown", ".csv": "CSV"}


class NoSePudoLeer(Exception):
    """El archivo no dio texto. Se distingue de un error del sistema porque
    el usuario puede hacer algo al respecto."""


# =====================================================================
# EXTRACCIÓN
# =====================================================================

def _pdf(ruta: Path) -> list[tuple[int, str]]:
    from pypdf import PdfReader
    lector = PdfReader(str(ruta))
    paginas = []
    for i, pag in enumerate(lector.pages, start=1):
        texto = (pag.extract_text() or "").strip()
        if texto:
            paginas.append((i, texto))
    if not paginas:
        raise NoSePudoLeer(
            "El PDF no tiene texto extraíble: probablemente es un escaneo. "
            "Hay que pasarle OCR antes de subirlo.")
    return paginas


def _docx(ruta: Path) -> list[tuple[int, str]]:
    import docx
    d = docx.Document(str(ruta))
    partes = [p.text for p in d.paragraphs if p.text.strip()]
    # Las tablas de Word traen buena parte del contenido en documentos de
    # trabajo, y `paragraphs` no las incluye.
    for t in d.tables:
        for fila in t.rows:
            celdas = [c.text.strip() for c in fila.cells if c.text.strip()]
            if celdas:
                partes.append(" | ".join(celdas))
    if not partes:
        raise NoSePudoLeer("El documento de Word no tiene texto.")
    return [(None, "\n".join(partes))]


def _xlsx(ruta: Path) -> list[tuple[int, str]]:
    from openpyxl import load_workbook
    wb = load_workbook(str(ruta), read_only=True, data_only=True)
    bloques = []
    for hoja in wb.worksheets:
        filas = []
        for fila in hoja.iter_rows(values_only=True):
            celdas = [str(c) for c in fila if c is not None and str(c).strip()]
            if celdas:
                filas.append(" | ".join(celdas))
        if filas:
            bloques.append((None, f"Hoja: {hoja.title}\n" + "\n".join(filas)))
    if not bloques:
        raise NoSePudoLeer("El libro de Excel no tiene celdas con contenido.")
    return bloques


def _texto(ruta: Path) -> list[tuple[int, str]]:
    for cod in ("utf-8", "cp1252", "latin-1"):
        try:
            t = ruta.read_text(encoding=cod).strip()
            if t:
                return [(None, t)]
            raise NoSePudoLeer("El archivo está vacío.")
        except UnicodeDecodeError:
            continue
    raise NoSePudoLeer("No se pudo determinar la codificación del archivo.")


LECTORES = {"PDF": _pdf, "Word": _docx, "Excel": _xlsx,
            "Texto": _texto, "Markdown": _texto, "CSV": _texto}


def extraer(ruta: Path, tipo: str) -> list[tuple[int | None, str]]:
    lector = LECTORES.get(tipo)
    if not lector:
        raise NoSePudoLeer(f"No sé leer archivos de tipo {tipo}.")
    return lector(ruta)


# =====================================================================
# TROCEADO
# =====================================================================

def trocear(bloques: list[tuple[int | None, str]],
            tipo: str | None = None) -> list[dict]:
    """Corta el texto en fragmentos que se entiendan solos.

    Se corta en el límite de párrafo más cercano y no en el carácter exacto:
    partir una frase por la mitad produce dos fragmentos que no significan
    nada y que el buscador nunca va a encontrar bien.

    El solape existe porque la respuesta a una pregunta suele estar a
    caballo entre dos fragmentos; repetir el final del anterior al principio
    del siguiente evita perderla justo en la costura.
    """
    tope = TAMANO_TABLA if tipo in TIPOS_TABLA else TAMANO
    fragmentos: list[dict] = []
    for pagina, texto in bloques:
        texto = re.sub(r"[ \t]+", " ", texto).strip()
        parrafos = [p.strip() for p in re.split(r"\n\s*\n", texto) if p.strip()]
        if not parrafos:
            continue

        actual = ""
        for p in parrafos:
            if len(actual) + len(p) + 2 <= tope:
                actual = f"{actual}\n\n{p}" if actual else p
                continue
            if actual:
                fragmentos.append({"pagina": pagina, "texto": actual})
                actual = actual[-SOLAPE:] if SOLAPE else ""
            # Un párrafo más largo que el tope se corta por longitud: no hay
            # límite natural donde partirlo.
            while len(p) > tope:
                # `rfind` devuelve -1 cuando no hay espacio, y -1 es verdadero
                # en Python: un `or` aquí dejaría pasar el párrafo entero.
                corte = p.rfind(" ", 0, tope)
                if corte <= 0:
                    corte = tope
                fragmentos.append({"pagina": pagina,
                                   "texto": (actual + " " + p[:corte]).strip()})
                p, actual = p[corte:].strip(), ""
            actual = f"{actual}\n\n{p}".strip() if actual else p
        if actual.strip():
            fragmentos.append({"pagina": pagina, "texto": actual.strip()})

    for i, f in enumerate(fragmentos):
        f["orden"] = i
    return fragmentos


# =====================================================================
# PROCESO COMPLETO
# =====================================================================

def huella(datos: bytes) -> str:
    return hashlib.sha256(datos).hexdigest()


def guardar_archivo(datos: bytes, nombre: str) -> tuple[Path, str]:
    """Escribe el archivo con su huella por nombre. Dos usuarios que suban
    el mismo archivo comparten el fichero en disco; el registro en la base
    sigue siendo de cada quien."""
    ARCHIVOS.mkdir(parents=True, exist_ok=True)
    h = huella(datos)
    destino = ARCHIVOS / f"{h}{Path(nombre).suffix.lower()}"
    if not destino.exists():
        destino.write_bytes(datos)
    return destino, h


def procesar(doc_id: str) -> dict:
    """Extrae, trocea y vectoriza. Deja el documento LISTO o en ERROR, nunca
    a medias sin decirlo."""
    d = db.documento(doc_id)
    if not d:
        raise ValueError(f"El documento {doc_id} no existe")

    try:
        db.actualizar_documento(doc_id, estado="EXTRAYENDO", error=None)

        # Una hoja de cálculo se guarda ADEMÁS con su forma de tabla. El
        # texto sirve para encontrarla por parecido; la tabla sirve para
        # contar, listar y sumar sin muestrear -- que es lo que un modelo
        # nunca va a poder hacer sobre fragmentos.
        if d["tipo"] in ("Excel", "CSV"):
            tabla.guardar(doc_id, tabla.leer(Path(d["archivo"]), d["tipo"]))

        bloques = extraer(Path(d["archivo"]), d["tipo"])
        caracteres = sum(len(t) for _, t in bloques)
        paginas = sum(1 for p, _ in bloques if p is not None) or None

        fragmentos = trocear(bloques, d["tipo"])
        if not fragmentos:
            raise NoSePudoLeer("El archivo no produjo ningún fragmento de texto.")

        db.actualizar_documento(doc_id, estado="VECTORIZANDO",
                                caracteres=caracteres, paginas=paginas)

        # Por lotes: vectorizar 400 fragmentos de un golpe en CPU agota la
        # memoria del proceso de Ollama.
        LOTE = int(os.getenv("EMBED_LOTE", "16"))
        for i in range(0, len(fragmentos), LOTE):
            trozo = fragmentos[i:i + LOTE]
            vectores = ia.embeddings([f["texto"] for f in trozo])
            for f, v in zip(trozo, vectores):
                f["embedding"] = v
                f["modelo_embed"] = ia.MODELO_EMBED

        n = db.guardar_fragmentos(doc_id, fragmentos)
        db.actualizar_documento(doc_id, estado="LISTO", n_fragmentos=n,
                                procesado_en=datetime.now(timezone.utc),
                                error=None)
        return {"estado": "LISTO", "fragmentos": n, "caracteres": caracteres,
                "paginas": paginas}

    except NoSePudoLeer as e:
        db.actualizar_documento(doc_id, estado="ERROR", error=str(e))
        return {"estado": "ERROR", "error": str(e)}
    except Exception as e:
        # El detalle técnico se guarda tal cual: adivinar qué falló y
        # traducirlo a algo bonito esconde justo lo que hace falta para
        # arreglarlo.
        db.actualizar_documento(doc_id, estado="ERROR",
                                error=f"{type(e).__name__}: {e}")
        return {"estado": "ERROR", "error": f"{type(e).__name__}: {e}"}
